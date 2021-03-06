import docx, os, sys, getopt

filename = "AssemledDocs.docx"

def getAllDocx(directory = os.getcwd()):
    files = []
    try:
        for f in os.listdir(directory):
            if f.endswith('.docx') or f.endswith('.doc'):
                files.append(f)
    except:
        print(f"В директории {directory} нет документов")
        sys.exit()
    return files

# start recording text on "ЭКЗАМЕНАЦИОННЫЙ БИЛЕТ № **" end on "Составитель ст.преподаватель кафедры «РЭТ»                           Наурыз К.Ж."
def copyAllDocsToOne(path, mainDoc, docs, startPhrase, endPhrase):
    if len(docs) == 0:
        print(f"Нет документов в директории {path}")
        printHelp()
        sys.exit()
    for docName in docs:
        copyFlag = False
        paragraphsToCopy = []
        try:
            partOfEverythingDoc = docx.Document(os.path.join(path, docName))
        except:
            print("В этой директории нет документов")
            sys.exit()
        for paragraph in partOfEverythingDoc.paragraphs:
            if startPhrase in paragraph.text:
                copyFlag = True
            elif endPhrase in paragraph.text:
                copyFlag = False
            if copyFlag:
                paragraphsToCopy.append(paragraph.text)
        addParagraphsToMainDoc(mainDoc, paragraphsToCopy)
        print(f"{docName} скопирован")
    if len(docs) != 0:
        mainDoc.save(filename)
        print(f"Результат: {os.path.join(os.getcwd(), filename)}")

def addParagraphsToMainDoc(mainDoc, paragraphsToCopy):
    for i in paragraphsToCopy:
        mainDoc.add_paragraph(i)

def printHelp():
    print('docsAssembler.py p <path> s <startphrase> -e <endphrase>')

def main(argv):
    # path = ''
    path = '/home/deus/a'
    # startPhrase = ''
    startPhrase = "ЭКЗАМЕНАЦИОННЫЙ БИЛЕТ"
    # endPhrase = ''
    endPhrase = "Составитель ст.преподаватель кафедры «РЭТ»"
    try:
        opts, _ = getopt.getopt(argv,"hp:s:e:",["path=","startph=","endph="])
    except getopt.GetoptError:
        printHelp()
        sys.exit(2)
    for opt, arg in opts:
        if opt == '-h':
            printHelp()
            sys.exit()
        elif opt in ("-s", "--startph"):
            startPhrase = arg
        elif opt in ("-e", "--endph"):
            endPhrase = arg
        elif opt in ("-p", "--path"):
            path = arg
    if path == "":
        path = os.getcwd()
    docs = getAllDocx(path)
    mainDoc = docx.Document()
    copyAllDocsToOne(path, mainDoc, docs, startPhrase, endPhrase)

if __name__ == "__main__":
   main(sys.argv[1:])
